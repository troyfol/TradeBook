"""Multi-format document exporter for journal entries and briefs.

Supported formats:
    * ``docx``      — Word document with embedded images (python-docx)
    * ``txt``       — Plain text; images stripped
    * ``md``        — Markdown; images saved next to the file in a
                      ``<name>_attachments/`` folder and referenced
                      relatively
    * ``html``      — Self-contained HTML with images base64-embedded
                      as ``data:`` URIs (one portable file)

Input is the same HTML format used by ``journal_entries`` and ``briefs``:
a Qt-authored rich-text document where images are referenced as
``attachment:N`` URLs resolving to rows in ``journal_attachments``.

The exporter resolves attachments through a supplied ``fetch_attachment``
callback (so tests can inject fakes) and never touches the UI layer.
"""
from __future__ import annotations

import base64
import mimetypes
import os
import re
import sqlite3
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Callable, Optional

from ingest import db_manager


FORMAT_DOCX = "docx"
FORMAT_TXT = "txt"
FORMAT_MARKDOWN = "md"
FORMAT_HTML = "html"

EXPORT_FORMATS: list[tuple[str, str]] = [
    (FORMAT_DOCX, "Word document (.docx)"),
    (FORMAT_TXT, "Plain text (.txt)"),
    (FORMAT_MARKDOWN, "Markdown (.md)"),
    (FORMAT_HTML, "HTML (.html)"),
]


# ---- block model ----------------------------------------------------------


@dataclass
class Block:
    """One renderable unit extracted from the source HTML."""
    kind: str  # "h1" | "h2" | "h3" | "p" | "image" | "hr" | "list_item"
    text: str = ""
    # For images: the attachment id (or ``None`` if the src couldn't be parsed)
    attachment_id: Optional[int] = None


@dataclass
class Document:
    title: str
    blocks: list[Block] = field(default_factory=list)


# ---- HTML → Block list ----------------------------------------------------


_ATTACHMENT_URL_RE = re.compile(r"^attachment:(\d+)", re.IGNORECASE)

_HEADING_TAGS = {"h1", "h2", "h3", "h4", "h5", "h6"}
_PARAGRAPH_TAGS = {"p", "div"}
_LIST_ITEM_TAGS = {"li"}
# Tags to discard entirely along with their contents.
_SKIP_TAGS = {"script", "style", "head", "title"}
# Void elements never emit an end tag. Counting them while unwinding a
# skipped region would leave ``_skip_depth`` permanently above zero and
# silently swallow the rest of the document — a bare `<meta charset>`
# inside <head> was enough to export an empty file. (Qt self-closes its
# tags, which routes them through ``handle_startendtag`` and hid this.)
_VOID_TAGS = {
    "area", "base", "br", "col", "embed", "hr", "img", "input",
    "link", "meta", "param", "source", "track", "wbr",
}


class _HtmlToBlocks(HTMLParser):
    """Walk Qt-emitted HTML and emit a flat list of ``Block``.

    Handles headings, paragraphs, list items, horizontal rules, line
    breaks, and ``<img src='attachment:N'>`` references. Ignores styling
    attributes and inline formatting (bold/italic/etc.) — the resulting
    plain text is fine for txt/md/docx which all re-wrap paragraphs on
    their own.
    """

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[Block] = []
        self._current_kind: Optional[str] = None
        self._current_text: list[str] = []
        self._skip_depth = 0

    # ---- buffer management ------------------------------------------------

    def _flush(self) -> None:
        if self._current_kind is None:
            # Free-floating text outside any block — wrap it as a paragraph.
            text = "".join(self._current_text).strip()
            if text:
                self.blocks.append(Block(kind="p", text=text))
        else:
            text = "".join(self._current_text).strip()
            if text or self._current_kind == "hr":
                self.blocks.append(Block(kind=self._current_kind, text=text))
        self._current_kind = None
        self._current_text = []

    def _start_block(self, kind: str) -> None:
        self._flush()
        self._current_kind = kind
        self._current_text = []

    # ---- parser hooks -----------------------------------------------------

    def handle_starttag(self, tag: str, attrs) -> None:
        if self._skip_depth:
            if tag not in _VOID_TAGS:
                self._skip_depth += 1
            return
        if tag in _SKIP_TAGS:
            self._skip_depth = 1
            return
        if tag == "img":
            for name, val in attrs:
                if name == "src" and val:
                    m = _ATTACHMENT_URL_RE.match(val.strip())
                    if m:
                        self._flush()
                        self.blocks.append(Block(
                            kind="image",
                            attachment_id=int(m.group(1)),
                        ))
                        return
            # Non-attachment image — silently ignore (we don't embed
            # external URLs).
            return
        if tag == "br":
            self._current_text.append("\n")
            return
        if tag == "hr":
            self._flush()
            self.blocks.append(Block(kind="hr"))
            return
        if tag in _HEADING_TAGS:
            self._start_block(tag)
            return
        if tag in _LIST_ITEM_TAGS:
            self._start_block("list_item")
            return
        if tag in _PARAGRAPH_TAGS:
            self._start_block("p")
            return
        # Other tags (spans, bold, etc.) — ignore markup, keep text.

    def handle_endtag(self, tag: str) -> None:
        if self._skip_depth:
            self._skip_depth -= 1
            return
        if tag in _HEADING_TAGS or tag in _LIST_ITEM_TAGS or tag in _PARAGRAPH_TAGS:
            self._flush()

    def handle_data(self, data: str) -> None:
        if self._skip_depth:
            return
        self._current_text.append(data)

    def close(self) -> None:
        super().close()
        self._flush()


def html_to_blocks(html: str) -> list[Block]:
    p = _HtmlToBlocks()
    p.feed(html or "")
    p.close()
    # Collapse consecutive empty blocks; trim trailing whitespace-only.
    out: list[Block] = []
    for b in p.blocks:
        if b.kind == "p" and not b.text.strip():
            continue
        out.append(b)
    return out


# ---- attachment resolution ------------------------------------------------


AttachmentFetcher = Callable[[int], Optional[tuple[str, bytes, Optional[str]]]]


def _safe_basename(name: str) -> str:
    name = os.path.basename(name or "")
    return re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", name)


def _extension_for(mime: str, filename: Optional[str]) -> str:
    if filename:
        ext = os.path.splitext(filename)[1]
        if ext:
            return ext
    return mimetypes.guess_extension(mime or "") or ".bin"


# ---- TXT ------------------------------------------------------------------


def _write_txt(doc: Document, out_path: Path) -> None:
    lines: list[str] = [doc.title, "=" * max(3, len(doc.title)), ""]
    for b in doc.blocks:
        if b.kind == "hr":
            lines.append("-" * 40)
            lines.append("")
        elif b.kind == "image":
            # Plain text has no image rep — leave a placeholder so the
            # structure stays readable.
            lines.append("[image]")
            lines.append("")
        elif b.kind in _HEADING_TAGS:
            lines.append(b.text)
            lines.append("-" * max(3, len(b.text)))
            lines.append("")
        elif b.kind == "list_item":
            lines.append(f"  - {b.text}")
        else:  # p
            lines.append(b.text)
            lines.append("")
    out_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


# ---- Markdown -------------------------------------------------------------


def _write_markdown(
    doc: Document,
    out_path: Path,
    fetch: AttachmentFetcher,
) -> None:
    # Images go into a sibling folder so relative links work for any
    # markdown renderer.
    stem = out_path.stem
    attach_dir = out_path.parent / f"{stem}_attachments"
    attach_dir_created = False

    lines: list[str] = [f"# {doc.title}", ""]
    img_counter = 0
    for b in doc.blocks:
        if b.kind == "h1":
            lines.append(f"# {b.text}")
            lines.append("")
        elif b.kind == "h2":
            lines.append(f"## {b.text}")
            lines.append("")
        elif b.kind == "h3":
            lines.append(f"### {b.text}")
            lines.append("")
        elif b.kind in _HEADING_TAGS:  # h4-h6
            lines.append(f"#### {b.text}")
            lines.append("")
        elif b.kind == "hr":
            lines.append("---")
            lines.append("")
        elif b.kind == "list_item":
            lines.append(f"- {b.text}")
        elif b.kind == "image":
            if b.attachment_id is None:
                continue
            payload = fetch(b.attachment_id)
            if payload is None:
                continue
            mime, data, filename = payload
            img_counter += 1
            ext = _extension_for(mime, filename)
            safe = _safe_basename(filename) if filename else ""
            img_name = (
                f"img_{img_counter:03d}_{safe}"
                if safe else f"img_{img_counter:03d}{ext}"
            )
            if not attach_dir_created:
                attach_dir.mkdir(parents=True, exist_ok=True)
                attach_dir_created = True
            (attach_dir / img_name).write_bytes(data)
            rel = f"./{attach_dir.name}/{img_name}"
            lines.append(f"![]({rel})")
            lines.append("")
        else:  # p
            lines.append(b.text)
            lines.append("")

    out_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


# ---- HTML -----------------------------------------------------------------


def _write_html(
    doc: Document,
    out_path: Path,
    fetch: AttachmentFetcher,
) -> None:
    # Self-contained HTML: inline every attachment as a base64 data URI
    # so the file is portable with no sibling folder required.
    from html import escape as _esc

    parts: list[str] = [
        "<!doctype html>",
        "<html><head>",
        '<meta charset="utf-8">',
        f"<title>{_esc(doc.title)}</title>",
        "<style>",
        "body{font-family:system-ui,-apple-system,'Segoe UI',sans-serif;"
        "max-width:780px;margin:2em auto;padding:0 1em;line-height:1.5;}",
        "img{max-width:100%;height:auto;display:block;margin:1em auto;}",
        "hr{border:0;border-top:1px solid #ccc;margin:2em 0;}",
        "h1,h2,h3{color:#222;}",
        "</style></head><body>",
        f"<h1>{_esc(doc.title)}</h1>",
    ]
    for b in doc.blocks:
        if b.kind == "hr":
            parts.append("<hr />")
        elif b.kind == "image":
            if b.attachment_id is None:
                continue
            payload = fetch(b.attachment_id)
            if payload is None:
                continue
            mime, data, _filename = payload
            b64 = base64.b64encode(data).decode("ascii")
            parts.append(
                f'<img src="data:{mime};base64,{b64}" alt="" />'
            )
        elif b.kind in _HEADING_TAGS:
            parts.append(f"<{b.kind}>{_esc(b.text)}</{b.kind}>")
        elif b.kind == "list_item":
            parts.append(f"<li>{_esc(b.text)}</li>")
        else:  # p
            parts.append(f"<p>{_esc(b.text)}</p>")
    parts.append("</body></html>")
    out_path.write_text("\n".join(parts), encoding="utf-8")


# ---- DOCX -----------------------------------------------------------------


def _write_docx(
    doc: Document,
    out_path: Path,
    fetch: AttachmentFetcher,
) -> None:
    # Local import so docx isn't needed for txt/md/html exports.
    from docx import Document as DocxDocument
    from docx.shared import Inches
    from io import BytesIO

    d = DocxDocument()
    d.add_heading(doc.title, level=0)

    for b in doc.blocks:
        if b.kind == "h1":
            d.add_heading(b.text, level=1)
        elif b.kind == "h2":
            d.add_heading(b.text, level=2)
        elif b.kind == "h3":
            d.add_heading(b.text, level=3)
        elif b.kind in _HEADING_TAGS:
            d.add_heading(b.text, level=4)
        elif b.kind == "hr":
            # python-docx has no horizontal rule primitive; use an empty
            # paragraph as a visual break instead.
            d.add_paragraph("———————————")
        elif b.kind == "list_item":
            d.add_paragraph(b.text, style="List Bullet")
        elif b.kind == "image":
            if b.attachment_id is None:
                continue
            payload = fetch(b.attachment_id)
            if payload is None:
                continue
            _mime, data, _filename = payload
            try:
                d.add_picture(BytesIO(data), width=Inches(5.0))
            except Exception:
                # Unsupported image format — skip rather than abort.
                continue
        elif b.text:
            d.add_paragraph(b.text)

    d.save(str(out_path))


# ---- public entry ---------------------------------------------------------


def export_document(
    html: str,
    title: str,
    fmt: str,
    out_path: str | Path,
    *,
    fetch_attachment: AttachmentFetcher,
    include_images: bool = True,
) -> Path:
    """Export a document (journal entry or brief) to disk.

    Parameters
    ----------
    html : str
        Source HTML (Qt-emitted rich text with ``attachment:N`` image
        refs).
    title : str
        The document title — rendered as the top-level heading.
    fmt : str
        One of ``FORMAT_DOCX / FORMAT_TXT / FORMAT_MARKDOWN / FORMAT_HTML``.
    out_path : str | Path
        Full output path. The extension is NOT auto-added; the caller is
        responsible for building the right filename.
    fetch_attachment : callable
        ``(attachment_id) -> (mime, bytes, filename) | None``. Typically
        ``lambda aid: db_manager.fetch_attachment(conn, aid)``.
    include_images : bool
        If False, image blocks are dropped before rendering (txt is
        always imageless regardless of this flag).
    """
    blocks = html_to_blocks(html)
    if not include_images:
        blocks = [b for b in blocks if b.kind != "image"]
    doc = Document(title=title, blocks=blocks)
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    if fmt == FORMAT_TXT:
        _write_txt(doc, out_path)
    elif fmt == FORMAT_MARKDOWN:
        _write_markdown(doc, out_path, fetch_attachment)
    elif fmt == FORMAT_HTML:
        _write_html(doc, out_path, fetch_attachment)
    elif fmt == FORMAT_DOCX:
        _write_docx(doc, out_path, fetch_attachment)
    else:
        raise ValueError(f"Unknown export format: {fmt!r}")
    return out_path


def extension_for_format(fmt: str) -> str:
    return {
        FORMAT_DOCX: ".docx",
        FORMAT_TXT: ".txt",
        FORMAT_MARKDOWN: ".md",
        FORMAT_HTML: ".html",
    }[fmt]


def make_attachment_fetcher(
    conn: sqlite3.Connection,
) -> AttachmentFetcher:
    """Standard attachment lookup used by the UI layer."""
    def _fetch(att_id: int):
        return db_manager.fetch_attachment(conn, att_id)
    return _fetch
